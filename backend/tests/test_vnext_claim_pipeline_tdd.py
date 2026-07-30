from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document

from backend.vnext.artifacts.local_store import LocalArtifactStore
from backend.vnext.claims import (
    atomize_source_claims,
    audit_claim_omissions,
    evaluate_omission_audit,
)
from backend.vnext.contracts.claims import (
    ClaimPublicationStatus,
    ClaimType,
    ExternalValidityStatus,
    SourceEntailmentStatus,
)
from backend.vnext.contracts.common import (
    ArtifactProducerRef,
    ArtifactRef,
    ArtifactType,
    RuntimeRole,
)
from backend.vnext.orchestration.source_shadow import run_source_shadow

from backend.tests.vnext_test_support import artifact_id, digest, region_id


def _region_ref(owner: str = "tenant-a") -> ArtifactRef:
    return ArtifactRef(
        owner_id=owner,
        artifact_id=artifact_id("9"),
        artifact_type=ArtifactType.REGION_PLAN,
        payload_digest=digest("9"),
    )


def _mapping(result, *, excluded: set[str] | None = None) -> dict[str, str]:
    excluded = excluded or set()
    return {
        entry.source_id: region_id("1")
        for entry in result.source_inventory.all_entries()
        if entry.source_id not in excluded
    }


class VNextClaimPipelineTests(unittest.TestCase):
    def test_source_only_claims_keep_instruction_and_split_states(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            source_path = root / "course.md"
            source_path.write_text(
                "# Carbonyl Compounds\n"
                "Aldehydes are compounds with a terminal carbonyl group.\n"
                "Complete the following conversion.\n"
                "E = mc^2\n",
                encoding="utf-8",
            )
            store = LocalArtifactStore(root / "shadow")
            source_result = run_source_shadow(
                source_path,
                owner_id="tenant-a",
                store=store,
            )
            region_ref = _region_ref()

            ledger = atomize_source_claims(
                source_result.source_observation,
                document_ir_ref=store.ref(
                    source_result.source_envelope
                ),
                region_plan_refs=(region_ref,),
                source_to_leaf_region=_mapping(source_result),
            )
            repeated = atomize_source_claims(
                source_result.source_observation,
                document_ir_ref=store.ref(
                    source_result.source_envelope
                ),
                region_plan_refs=(region_ref,),
                source_to_leaf_region=_mapping(source_result),
            )

            self.assertEqual(ledger, repeated)
            self.assertTrue(ledger.claims)
            instruction = next(
                claim
                for claim in ledger.claims
                if claim.claim_type is ClaimType.INSTRUCTION
            )
            self.assertEqual(
                instruction.publication_status,
                ClaimPublicationStatus.WITHHELD,
            )
            self.assertEqual(
                instruction.source_entailment_status,
                SourceEntailmentStatus.ENTAILED,
            )
            self.assertEqual(
                instruction.external_validity_status,
                ExternalValidityStatus.NOT_CHECKED,
            )
            self.assertEqual(instruction.external_evidence_refs, ())
            self.assertIsNotNone(instruction.fidelity_verifier)
            assert instruction.fidelity_verifier is not None
            self.assertNotEqual(
                instruction.extractor.producer_id,
                instruction.fidelity_verifier.producer_id,
            )
            self.assertTrue(
                any(
                    claim.publication_status
                    is ClaimPublicationStatus.CORE
                    for claim in ledger.claims
                )
            )

    def test_independent_omission_audit_reconciles_complete_ledger(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            source_path = root / "course.md"
            source_path.write_text(
                "# Carbonyl Compounds\n"
                "Aldehydes are compounds with a terminal carbonyl group.\n"
                "Complete the following conversion.\n",
                encoding="utf-8",
            )
            store = LocalArtifactStore(root / "shadow")
            source_result = run_source_shadow(
                source_path,
                owner_id="tenant-a",
                store=store,
            )
            source_ref = store.ref(source_result.source_envelope)
            inventory_ref = store.ref(source_result.inventory_envelope)
            region_ref = _region_ref()
            ledger = atomize_source_claims(
                source_result.source_observation,
                document_ir_ref=source_ref,
                region_plan_refs=(region_ref,),
                source_to_leaf_region=_mapping(source_result),
            )
            ledger_envelope = store.put(
                owner_id="tenant-a",
                role=RuntimeRole.CLAIM_ATOMIZER,
                payload=ledger,
                producer=ArtifactProducerRef(
                    producer_id="vnext-source-claim-atomizer",
                    producer_version="1.0.0",
                    role=RuntimeRole.CLAIM_ATOMIZER,
                ),
                input_refs=(source_ref, region_ref),
            )
            ledger_ref = store.ref(ledger_envelope)

            audit = audit_claim_omissions(
                source_result.source_inventory,
                ledger,
                source_inventory_ref=inventory_ref,
                claim_ledger_ref=ledger_ref,
            )

            self.assertEqual(audit.omitted_source_ids, ())
            self.assertEqual(
                audit.high_importance_omitted_source_ids,
                (),
            )
            self.assertTrue(
                evaluate_omission_audit(
                    source_result.source_inventory,
                    audit,
                ).accepted
            )
            all_inventory_ids = {
                entry.source_id
                for entry in source_result.source_inventory.all_entries()
            }
            reconciled = {
                *audit.accounted_source_ids,
                *audit.unresolved_source_ids,
                *audit.explicitly_nonclaim_source_ids,
            }
            self.assertEqual(reconciled, all_inventory_ids)

    def test_missing_high_value_source_is_not_hidden(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            source_path = root / "course.md"
            source_path.write_text(
                "# Carbonyl Compounds\n"
                "Aldehydes are compounds with a terminal carbonyl group.\n",
                encoding="utf-8",
            )
            store = LocalArtifactStore(root / "shadow")
            source_result = run_source_shadow(
                source_path,
                owner_id="tenant-a",
                store=store,
            )
            source_ref = store.ref(source_result.source_envelope)
            inventory_ref = store.ref(source_result.inventory_envelope)
            region_ref = _region_ref()
            heading_id = next(
                entry.source_id
                for entry in source_result.source_inventory.block_entries
                if entry.declared_role == "heading"
            )
            ledger = atomize_source_claims(
                source_result.source_observation,
                document_ir_ref=source_ref,
                region_plan_refs=(region_ref,),
                source_to_leaf_region=_mapping(
                    source_result,
                    excluded={heading_id},
                ),
            )
            ledger_envelope = store.put(
                owner_id="tenant-a",
                role=RuntimeRole.CLAIM_ATOMIZER,
                payload=ledger,
                producer=ArtifactProducerRef(
                    producer_id="vnext-source-claim-atomizer",
                    producer_version="1.0.0",
                    role=RuntimeRole.CLAIM_ATOMIZER,
                ),
                input_refs=(source_ref, region_ref),
            )
            audit = audit_claim_omissions(
                source_result.source_inventory,
                ledger,
                source_inventory_ref=inventory_ref,
                claim_ledger_ref=store.ref(ledger_envelope),
            )

            self.assertIn(heading_id, audit.omitted_source_ids)
            self.assertIn(
                heading_id,
                audit.high_importance_omitted_source_ids,
            )
            gate = evaluate_omission_audit(
                source_result.source_inventory,
                audit,
            )
            self.assertFalse(gate.accepted)
            self.assertIn(
                "high_importance_source_omitted",
                gate.reason_codes,
            )

    def test_table_cells_are_individually_accounted(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            source_path = root / "course.docx"
            document = Document()
            document.add_heading("Carbonyl Comparison", level=1)
            table = document.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "Class"
            table.cell(0, 1).text = "Position"
            table.cell(1, 0).text = "Ketone"
            table.cell(1, 1).text = "Internal"
            document.save(source_path)
            store = LocalArtifactStore(root / "shadow")
            source_result = run_source_shadow(
                source_path,
                owner_id="tenant-a",
                store=store,
            )
            region_ref = _region_ref()
            ledger = atomize_source_claims(
                source_result.source_observation,
                document_ir_ref=store.ref(
                    source_result.source_envelope
                ),
                region_plan_refs=(region_ref,),
                source_to_leaf_region=_mapping(source_result),
            )
            claim_evidence_ids = {
                evidence.ref_id
                for claim in ledger.claims
                for evidence in claim.source_evidence_refs
            }
            table_cell_ids = {
                entry.source_id
                for entry in (
                    source_result.source_inventory.table_cell_entries
                )
            }

            self.assertEqual(
                table_cell_ids & claim_evidence_ids,
                table_cell_ids,
            )


if __name__ == "__main__":
    unittest.main()
